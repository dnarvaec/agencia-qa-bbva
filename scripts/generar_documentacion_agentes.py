# -*- coding: utf-8 -*-
"""
Genera un documento Word profesional con la documentacion de los agentes IA
del proyecto Agencia QA, aplicando la identidad visual de NTT DATA.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NTT_RED = RGBColor(0x00, 0x3D, 0x7A)
NTT_DARK = RGBColor(0x2B, 0x2B, 0x2B)
NTT_GRAY = RGBColor(0x58, 0x58, 0x58)
NTT_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEX_RED = "003D7A"
HEX_LIGHT_GRAY = "F2F2F2"
HEX_MID_GRAY = "D9D9D9"
FONT_NAME = "Calibri"

OUTPUT_DIR = os.path.join("archivos", "Documentacion")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Documentacion_Agentes_IA_NTT_DATA.docx")


def set_cell_bg(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def set_cell_borders(cell, color=HEX_MID_GRAY, size="4"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcpr.append(borders)


def cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_borders(cell)


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.bold = True
    run.font.size = Pt(20 if level == 1 else (15 if level == 2 else 12.5))
    run.font.color.rgb = NTT_RED if level == 1 else NTT_DARK
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12" if level == 1 else "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), HEX_RED if level == 1 else HEX_MID_GRAY)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def para(doc, text, size=10.5, color=None, bold=False, italic=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else NTT_DARK
    run.font.bold = bold
    run.font.italic = italic
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)
    run.font.color.rgb = NTT_DARK
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)
    run.font.color.rgb = NTT_DARK
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, HEX_RED)
        cell_text(c, h, bold=True, color=NTT_WHITE, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        bg = HEX_LIGHT_GRAY if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            set_cell_bg(cells[c_idx], bg)
            cell_text(cells[c_idx], val, size=9.8)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def footer_brand(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NTT DATA | Documentacion de Agentes IA - Proyecto Agencia QA")
    run.font.name = FONT_NAME
    run.font.size = Pt(8.5)
    run.font.color.rgb = NTT_GRAY


def set_margins(doc):
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)


def add_agent_section(doc, a, idx):
    heading(doc, "%d. %s" % (idx, a["name"]), level=1)
    para(doc, "Archivo fuente: %s" % a["archivo"], size=9.5, color=NTT_GRAY, italic=True, space_after=10)

    heading(doc, "Descripcion general", level=3)
    para(doc, a["descripcion"])

    heading(doc, "Principio fundamental", level=3)
    para(doc, a["principio"], bold=True, color=NTT_RED)

    heading(doc, "Ficha tecnica", level=3)
    table(doc, ["Atributo", "Detalle"], [
        ["Herramientas (tools)", a["tools"]],
        ["Stack tecnologico", a["stack"]],
        ["Modos de ejecucion", a.get("modos", "N/A")],
    ], widths=[5, 12])

    heading(doc, "Flujo de trabajo principal", level=3)
    for step in a["flujo"]:
        numbered(doc, step)

    heading(doc, "Reglas e invariantes clave", level=3)
    for r in a["reglas"]:
        bullet(doc, r)

    if a.get("rutas"):
        heading(doc, "Rutas y estructura de archivos", level=3)
        for r in a["rutas"]:
            bullet(doc, r)

    doc.add_page_break()


# ----------------------------------------------------------------------------
# Datos de los agentes
# ----------------------------------------------------------------------------
AGENTS = []

AGENTS.append({
    "name": "Automatizar y Ejecutar",
    "archivo": "Automatizar y Ejecutar.agent.md",
    "descripcion": (
        "Gestor integral de automatizacion de pruebas que analiza proyectos existentes, "
        "realiza exploracion obligatoria en vivo de la interfaz de usuario (UI) mediante las "
        "herramientas MCP de Playwright, y genera o mantiene codigo de automatizacion de "
        "pruebas listo para produccion con trazabilidad completa."
    ),
    "principio": "Ejecucion en vivo primero. Sin suposiciones, sin codigo teorico.",
    "tools": "vscode, execute, read, agent, edit, search, web, browser, playwright/*, todo",
    "stack": "Web: Page Object Model + TypeScript + Playwright + Playwright HTML Reporting. API: API Object Pattern + TypeScript + Playwright APIRequestContext.",
    "modos": "exploration - execution - replay - debug",
    "flujo": [
        "Analizar el proyecto existente y ejecutar las pruebas si ya existen",
        "Si es un flujo nuevo, realizar exploracion en vivo con Playwright MCP documentando todo en exploration_docs",
        "Generar codigo unicamente a partir de datos de exploracion validados",
        "Ejecutar, corregir errores y reejecutar hasta alcanzar el 100% de exito (maximo 5 iteraciones)",
        "Entregar codigo probado junto con reporte automatizado (capturas de pantalla y video)",
    ],
    "reglas": [
        "Prohibido generar codigo sin exploracion en vivo previa mediante MCP Playwright",
        "Prohibido usar selectores o estructuras DOM supuestas o teoricas",
        "Prohibido reportar exito o entregar codigo cuando existan pruebas fallando (falsos positivos)",
        "Todo fallo debe clasificarse como code_issue, system_bug o user_input_needed",
        "El reporte final debe incluir capturas de pantalla y video de exito y de fallo",
        "Generacion automatica de reportes de bugs (system_bug) en archivos/Seguimiento/{STORY_ID}/",
    ],
    "rutas": [
        "Casos de prueba a automatizar: archivos/Casos de Prueba/{CP_ID}/{CP_ID}-test-cases.json",
        "Automatizacion Web: carpeta 'automatizacion web/' (tests, src/pages)",
        "Automatizacion API: carpeta 'automatizacion api/' (tests, src/apis)",
        "Documentacion de exploracion: automatizacion web/exploration_docs/{CP_ID}/",
    ],
})


AGENTS.append({
    "name": "AZURE Disenar Casos de Prueba",
    "archivo": "AZURE Disenar Casos de Prueba.agent.md",
    "descripcion": (
        "Agente responsable de crear Casos de Prueba a partir de Historias de Usuario, "
        "asegurando la cobertura del 100% de la descripcion funcional y los criterios de "
        "aceptacion, con trazabilidad completa hacia Azure DevOps (Test Plans o Work Items "
        "segun el licenciamiento disponible de la cuenta)."
    ),
    "principio": "No inventar: toda la informacion se fundamenta exclusivamente en los criterios de aceptacion y la descripcion funcional de la HU.",
    "tools": "vscode, execute, read, agent, edit, search, web, browser, azure-devops/*, todo",
    "stack": "Generacion de casos web, api y manual con priorizacion: automatizable (web/api) > manual",
    "modos": "Diseno de casos (local) - Carga a Azure DevOps Modo A (Test Plans) o Modo B (Work Items)",
    "flujo": [
        "Leer la Historia de Usuario local desde archivos/HUs/{HU_ID}/{HU_ID}-final.json",
        "Disenar casos de prueba Web (positivo, negativo, por rol, limites, estado del sistema)",
        "Disenar casos de prueba API (positivo, autenticacion, validacion, recursos, permisos)",
        "Disenar casos Manuales solo cuando el escenario no sea automatizable",
        "Construir el JSON estructurado con resumen de conteo (total_web, total_api, total_manual, total)",
        "Generar el Markdown profesional a partir del JSON (tablas, trazabilidad, notas de automatizacion)",
        "Guardar ambos archivos en archivos/Casos de Prueba/{CP_ID}/",
        "Bajo demanda: cargar los casos a Azure DevOps (Test Plans o Work Items) y vincularlos a la HU",
    ],
    "reglas": [
        "Unicamente dos archivos de salida por ejecucion: {CP_ID}-test-cases.json y {CP_ID}-test-cases.md",
        "Prohibido crear archivos adicionales fuera del directorio especificado",
        "Prioridad de automatizacion: web y api por encima de manual",
        "Cada paso debe iniciar siempre con el login usando usuario y contrasena validos",
        "Carga a Azure DevOps: un solo Test Case por turno de herramientas (regla no negociable en Modo A)",
        "Vinculacion obligatoria de cada Test Case a la HU mediante el link 'Tested By'",
        "Verificacion post-carga: reintento si la Descripcion o los Steps no persistieron en Azure DevOps",
    ],
    "rutas": [
        "Historia de Usuario de entrada: archivos/HUs/{HU_ID}/{HU_ID}-final.json",
        "Casos de prueba (salida): archivos/Casos de Prueba/{CP_ID}/{CP_ID}-test-cases.json y .md",
        "Confirmacion de carga a Azure DevOps: archivos/Casos de Prueba/{CP_ID}/{CP_ID}-azure-upload.json y .md",
    ],
})


AGENTS.append({
    "name": "AZURE Mejorar HU",
    "archivo": "AZURE Mejorar HU.agent.md",
    "descripcion": (
        "Agente especializado en crear, evaluar y mejorar Historias de Usuario, tanto "
        "leyendolas y refinandolas desde Azure DevOps como generandolas desde cero a "
        "partir de informacion funcional entregada en el prompt, aplicando el estandar "
        "INVEST y publicandolas en Azure DevOps cuando se solicite."
    ),
    "principio": "Producir una Historia de Usuario bien estructurada, comprensible y sin ambiguedades, lista para el agente de diseno de casos de prueba.",
    "tools": "vscode, execute, read, agent, edit, search, web, browser, azure-devops/*, todo",
    "stack": "Estandar INVEST (Independiente, Negociable, Valiosa, Estimable, Pequena, Testeable) con scoring iterativo 1-10",
    "modos": "Flujo A: leer y mejorar HU existente - Flujo B: generar HU nueva - Flujo C: publicar en Azure DevOps - Flujo D: vincular a Test Plan",
    "flujo": [
        "Detectar el flujo a ejecutar segun las senales del prompt del usuario (ID numerico, descripcion funcional, palabras clave de publicacion o vinculacion)",
        "Flujo A: obtener la HU de Azure DevOps, limpiar HTML y evaluar con un score inicial",
        "Iterar hasta 3 veces mejorando descripcion y criterios de aceptacion (sin fusionar ni eliminar criterios) hasta alcanzar score >= 7",
        "Flujo B: extraer rol, funcionalidad y beneficio del prompt y redactar la HU desde cero con minimo 5 y maximo 12 criterios",
        "Guardar los entregables {ID}-final.json y {ID}-final.md en archivos/HUs/{ID}/",
        "Flujo C: publicar la HU como Work Item 'User Story' en Azure DevOps y actualizar el ID real en el JSON",
        "Flujo D: construir la jerarquia de Test Suites (Epica > Feature > HU) y vincular la HU a un Test Plan",
    ],
    "reglas": [
        "Los criterios de aceptacion de Azure DevOps se copian literalmente, nunca se parafrasean ni se acortan",
        "El array final de criterios debe tener igual o mas elementos que el original",
        "Cada criterio debe indicar QUE se valida, nunca COMO se implementa",
        "Si el score >= 7 o la mejora es minima, se detiene la iteracion",
        "Sin licencia de Test Plans, el Flujo D se detiene informando el error 403 (no tiene modo alternativo)",
        "Sin azure_devops_id valido, el Flujo D exige ejecutar primero el Flujo C de publicacion",
    ],
    "rutas": [
        "Entregables locales: archivos/HUs/{HU_ID}/{HU_ID}-final.json y {HU_ID}-final.md",
        "Proyecto Azure DevOps configurado via variable de entorno AZURE_DEVOPS_PROJECT",
    ],
})


AGENTS.append({
    "name": "JIRA Diseno de Casos de Prueba",
    "archivo": "JIRA Diseno de Casos de Prueba.md",
    "descripcion": (
        "Agente especializado en diseno de casos de prueba a partir de Historias de "
        "Usuario, con generacion de una suite completa en formato local (JSON/Markdown) "
        "y carga posterior directa a QMetry (QTM4J) dentro de la instancia Jira Data "
        "Center del cliente, mediante un script Python reutilizable."
    ),
    "principio": "No inventar: toda la informacion se fundamenta exclusivamente en los criterios de aceptacion y la descripcion funcional de la HU.",
    "tools": "vscode, execute, read, agent, edit, search, web, browser, todo",
    "stack": "Generacion de casos web, api y manual + integracion QMetry via API REST /rest/qtm4j/ui/latest (script jira_uploader.py, requests + python-dotenv)",
    "modos": "Diseno de casos (local) - Carga a QMetry (invocacion explicita del usuario)",
    "flujo": [
        "Leer la Historia de Usuario local desde archivos/HUs/{HU_ID}/",
        "Disenar casos de prueba Web, API y Manual siguiendo las mismas categorias que el agente de Azure",
        "Construir el JSON estructurado con resumen de conteo y guardarlo junto al Markdown en archivos/Casos de Prueba/{CP_ID}/",
        "Bajo demanda explicita del usuario, ejecutar jira_uploader.py para subir los casos como Test Cases de QMetry",
        "Mapear cada campo del JSON a los campos nativos de QMetry (summary, description, precondition, steps)",
        "Presentar el resumen final con las claves QMetry creadas (ej. CORREOF-TC-101) y las filas fallidas si existen",
    ],
    "reglas": [
        "Unicamente dos archivos de salida por diseno: {CP_ID}-test-cases.json y {CP_ID}-test-cases.md",
        "Prioridad de automatizacion: web y api por encima de manual",
        "La subida a QMetry solo se activa cuando el usuario lo solicita explicitamente en un prompt independiente",
        "El script jira_uploader.py se invoca por CLI (no como import) y no modifica el JSON de origen",
        "Todos los casos del arreglo test_cases se suben sin filtrar por tipo de automatizacion",
        "Autenticacion exclusivamente por Basic Auth usando JIRA_USERNAME y JIRA_API_TOKEN del .env",
    ],
    "rutas": [
        "Historia de Usuario de entrada: archivos/HUs/{HU_ID}/",
        "Casos de prueba (salida): archivos/Casos de Prueba/{CP_ID}/{CP_ID}-test-cases.json y .md",
        "Script de integracion: jira_uploader.py (raiz del proyecto) + requirements.txt",
        "Variables de entorno: JIRA_URL, JIRA_USERNAME, JIRA_API_TOKEN, QMETRY_PROJECT_ID (.env)",
    ],
})


# ----------------------------------------------------------------------------
# Construccion del documento
# ----------------------------------------------------------------------------
def build_cover(doc):
    section = doc.sections[0]

    # Barra superior de color corporativo
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(" ")
    run.font.size = Pt(2)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NTT DATA")
    run.font.name = FONT_NAME
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = NTT_RED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documentacion Tecnica de Agentes de Inteligencia Artificial")
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = NTT_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Proyecto Agencia QA")
    run.font.name = FONT_NAME
    run.font.size = Pt(15)
    run.font.color.rgb = NTT_GRAY

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Analisis funcional y tecnico de los agentes definidos en .github/agents")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = NTT_GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run(datetime.now().strftime("Fecha de generacion: %d/%m/%Y"))
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)
    run.font.color.rgb = NTT_DARK

    doc.add_page_break()


def build_intro(doc):
    heading(doc, "Introduccion", level=1)
    para(doc, (
        "Este documento presenta el analisis funcional y tecnico de los agentes de "
        "Inteligencia Artificial implementados en el proyecto Agencia QA. Cada agente "
        "esta definido como un archivo de configuracion Markdown ubicado en la carpeta "
        "'.github/agents/', y actua como un rol especializado dentro del ciclo de vida de "
        "calidad de software: desde la redaccion de Historias de Usuario, el diseno de "
        "casos de prueba, hasta la automatizacion y ejecucion de pruebas end-to-end."
    ))
    para(doc, (
        "Los cuatro agentes documentados a continuacion trabajan de forma complementaria, "
        "formando una cadena de valor de Aseguramiento de Calidad (QA) potenciada por IA: "
        "mejora de Historias de Usuario, diseno de casos de prueba (para Azure DevOps y "
        "Jira/QMetry), y automatizacion y ejecucion de pruebas Web y API."
    ))

    heading(doc, "Alcance del documento", level=3)
    bullet(doc, "Identidad y proposito de cada agente")
    bullet(doc, "Herramientas (tools) habilitadas para su ejecucion")
    bullet(doc, "Stack tecnologico y patrones de diseno aplicados")
    bullet(doc, "Flujo de trabajo principal paso a paso")
    bullet(doc, "Reglas e invariantes de comportamiento no negociables")
    bullet(doc, "Rutas y estructura de archivos de entrada/salida")

    doc.add_page_break()


def build_overview_table(doc):
    heading(doc, "Resumen Ejecutivo — Agentes del Proyecto", level=1)
    para(doc, "La siguiente tabla resume el proposito principal y el dominio de actuacion de cada agente:")

    rows = [
        [a["name"], a["archivo"], a["descripcion"][:110] + "..."] for a in AGENTS
    ]
    table(doc, ["Agente", "Archivo Fuente", "Proposito Principal"], rows, widths=[4.5, 4.5, 8])

    heading(doc, "Cadena de valor QA impulsada por IA", level=3)
    numbered(doc, "AZURE Mejorar HU — redacta o refina Historias de Usuario bajo el estandar INVEST")
    numbered(doc, "AZURE Disenar Casos de Prueba / JIRA Diseno de Casos de Prueba — generan casos de prueba detallados (web, api, manual) a partir de la HU")
    numbered(doc, "Automatizar y Ejecutar — convierte los casos de prueba en automatizacion real, ejecutada y validada en vivo")

    doc.add_page_break()


def build_closing(doc):
    heading(doc, "Conclusiones", level=1)
    para(doc, (
        "Los cuatro agentes analizados conforman un ecosistema integral de IA aplicada a "
        "Quality Assurance, cubriendo el ciclo completo desde la definicion de requerimientos "
        "hasta la validacion automatizada de software. Todos comparten principios comunes de "
        "trazabilidad, prohibicion de informacion inventada ('no inventar') y generacion de "
        "evidencia verificable en cada entrega."
    ))
    para(doc, (
        "Se recomienda mantener este documento actualizado cada vez que se modifique alguno "
        "de los archivos '.agent.md' en la carpeta '.github/agents', para preservar la "
        "trazabilidad entre la configuracion viva de los agentes y su documentacion oficial."
    ))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NTT DATA — Confidencial | Uso interno del proyecto Agencia QA")
    run.font.name = FONT_NAME
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = NTT_GRAY


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    set_margins(doc)
    footer_brand(doc)

    # Estilo base del documento
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)

    build_cover(doc)
    build_intro(doc)
    build_overview_table(doc)

    heading(doc, "Detalle de los Agentes", level=1)
    para(doc, "A continuacion se presenta la ficha tecnica completa de cada uno de los 4 agentes identificados en el proyecto.", space_after=14)

    for idx, agent in enumerate(AGENTS, start=1):
        add_agent_section(doc, agent, idx)

    build_closing(doc)

    doc.save(OUTPUT_FILE)
    print("Documento generado en: %s" % OUTPUT_FILE)


if __name__ == "__main__":
    main()
